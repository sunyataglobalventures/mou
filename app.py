from flask import Flask, request, render_template, send_file
import os
from docx import Document
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore
import json

app = Flask(__name__)

# Initialize Firebase Admin SDK
firebase_creds = os.getenv("GOOGLE_APPLICATION_CREDENTIALS_JSON")
if firebase_creds:
    # Parse JSON string from environment variable
    creds_dict = json.loads(firebase_creds)
    cred = credentials.Certificate(creds_dict)
    firebase_admin.initialize_app(cred)
else:
    raise RuntimeError("Firebase credentials not found in environment variables.")

# Firestore database reference
db = firestore.client()

def insert_submission(data):
    # Create a reference to the 'mou' collection
    mou_ref = db.collection('mou')
    
    # Add a new document with form data
    mou_ref.add({
        'name': data["name"],
        'email': data["email"],
        'address': data["address"],
        'storename': data["storename"],
        'pswrd': data["pswrd"],
        'service': data["service"],
        'cost': data["cost"],
        'duration': data["duration"],
        'timestamp': firestore.SERVER_TIMESTAMP  # Add server-side timestamp
    })

def replace_text_in_run(run, key, value):
    if key in run.text:
        run.text = run.text.replace(key, value)
        run.font.bold = True

def replace_placeholders(doc, placeholders):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in placeholders.items():
                if key in run.text:
                    print(f"Replacing {key} with {value}")
                    replace_text_in_run(run, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in placeholders.items():
                            if key in run.text:
                                print(f"Replacing {key} with {value}")
                                replace_text_in_run(run, key, value)

def generate_word(template_path, placeholders, output_folder):
    # Load template
    doc = Document(template_path)
    replace_placeholders(doc, placeholders)

    # Extract placeholders for naming
    service_name = placeholders.get("[SERVICE]", "SERVICE").replace("/", "_").replace("\\", "_")
    store_name = placeholders.get("[STORENAME]", "STORE").replace("/", "_").replace("\\", "_")
    date = datetime.now().strftime("%Y-%m-%d")

    # Format file name
    file_name = f"MOU_{service_name}_{store_name}_{date}.docx"
    output_path = os.path.join(output_folder, file_name)

    # Save the document
    doc.save(output_path)
    return output_path


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Parse form data
        data = request.form.to_dict()

        # Retrieve single service from form
        service = request.form.get("service")

        # Create placeholders to replace in Word template
        placeholders = {
            "[NAME]": data.get("name"),
            "[EMAIL]": data.get("email"),
            "[ADDRESS]": data.get("address"),
            "[STORENAME]": data.get("storename"),
            "[PSWRD]": data.get("pswrd"),
            "[SERVICE]": service,  # Only one service is selected
            "[COST]": data.get("cost"),
            "[DUR]": data.get("duration"),
            "DATE": datetime.now().strftime("%d/%m/%Y"),
        }

        # Insert form data into Firebase Firestore
        try:
            insert_submission(data)
        except Exception as e:
            print(f"Error inserting data into Firestore: {e}")

        # Output folder
        output_folder = "output"
        template_path = "MOU_TEMP.docx"

        # Generate Word document
        word_path = generate_word(template_path, placeholders, output_folder)
        print(f"Word file path: {word_path}")

        # Serve the generated Word file for download
        if "download_word" in request.form:
            try:
                print("Serving Word file for download.")
                return send_file(word_path, as_attachment=True)
            except Exception as e:
                print(f"Error sending file: {e}")
                return "An error occurred while downloading the file.", 500

    return render_template("index.html")


if __name__ == "__main__":
    if not os.path.exists("output"):
        os.makedirs("output")  # Ensure output directory exists
    app.run(debug=True)
